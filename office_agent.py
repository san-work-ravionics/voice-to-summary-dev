"""A tool-use agent standing in for what a Microsoft 365 Copilot Enterprise
Office agent would do — read a meeting transcript and act on it inside Word
and Excel, rather than only replying with text. See
phase5-office-agent/README.md for why Claude substitutes for Copilot here.

Claude-only: local/mistral in llm_provider.py go through a plain
transformers text-generation pipeline() call with no tool-calling support,
so this agent loop can't run on them — that's a real, reported Phase 5
finding (agentic tool use needs a cloud model), not a gap being papered over.
"""
import os

import anthropic
from docx import Document
from openpyxl import Workbook, load_workbook

from llm_provider import CLAUDE_MODEL

_NO_AUTH_ERROR = (
    "Claude API authentication failed. No credentials found — set "
    "ANTHROPIC_API_KEY, or run `ant auth login`."
)

_client_instance = None


def _get_client():
    global _client_instance
    if _client_instance is None:
        _client_instance = anthropic.Anthropic()
    return _client_instance


TOOLS = [
    {
        "name": "write_minutes_docx",
        "description": (
            "Write meeting minutes as a Word document. Call this once you "
            "have drafted the Topic/Key Points/Decisions/Actions content "
            "from the transcript."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "topic": {"type": "string"},
                "key_points": {"type": "array", "items": {"type": "string"}},
                "decisions": {"type": "array", "items": {"type": "string"}},
                "actions": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Each bullet must start with 'Person A: ' or 'Person B: '.",
                },
            },
            "required": ["topic", "key_points", "decisions", "actions"],
        },
    },
    {
        "name": "write_action_tracker_xlsx",
        "description": (
            "Append this meeting's action items as rows to the running "
            "Excel action tracker."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "actions": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "owner": {"type": "string", "description": "'Person A' or 'Person B'"},
                            "action": {"type": "string"},
                        },
                        "required": ["owner", "action"],
                    },
                },
            },
            "required": ["actions"],
        },
    },
]

SYSTEM_PROMPT = (
    "You are an office productivity agent, standing in for what a "
    "Microsoft 365 Copilot Enterprise agent would do inside Word/Excel: "
    "read a meeting transcript and act on it by calling tools, rather than "
    "only replying with text. There are exactly two speakers in the "
    "transcript. Never use personal names, even if one appears — refer to "
    "the two speakers only as 'Person A' and 'Person B'.\n\n"
    "For every request:\n"
    "1. Call write_minutes_docx exactly once with a one-line topic, "
    "key points, decisions, and actions (each Actions bullet starting with "
    "'Person A: ' or 'Person B: ', only for things a speaker personally "
    "committed to — phrases like 'I will'/'I'll'/'let me'). Never attribute "
    "an action to a third party who is merely mentioned in conversation.\n"
    "2. Call write_action_tracker_xlsx exactly once with the same action "
    "items, structured as {owner, action} pairs.\n"
    "3. Only after both tools have been called, reply with a short "
    "plain-text confirmation — this final reply is also scored as the "
    "run's brief, so it must itself be Markdown with exactly these four "
    "sections, matching what you wrote to the tools:\n"
    "## Topic\n(one line)\n"
    "## Key Points\n(bullet points)\n"
    "## Decisions\n(bullet points)\n"
    "## Actions\n(bullet points, each 'Person A: ...' or 'Person B: ...')\n"
    "If a section has nothing to report, write a single bullet saying so."
)


def _write_minutes_docx(path, tool_input):
    doc = Document()
    doc.add_heading(tool_input.get("topic", ""), level=1)
    for section, key in (("Key Points", "key_points"), ("Decisions", "decisions"), ("Actions", "actions")):
        doc.add_heading(section, level=2)
        items = tool_input.get(key) or []
        if not items:
            doc.add_paragraph("Nothing to report.")
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.save(path)


def _append_action_tracker_xlsx(path, actions):
    if os.path.exists(path):
        wb = load_workbook(path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Action Tracker"
        ws.append(["Owner", "Action"])
    for item in actions or []:
        ws.append([item.get("owner", ""), item.get("action", "")])
    wb.save(path)


def run_agent(transcript, output_dir, model_name=None, max_turns=6, network_call=None):
    """Runs the tool-use loop over `transcript`, writing minutes.docx and
    action_tracker.xlsx into output_dir as Claude calls the corresponding
    tools. `network_call(fn)`, if given, wraps each actual API call — used
    by phase5-office-agent/src/network_sim.py to inject simulated
    latency/timeout/offline behavior around the network hop; defaults to
    calling fn() directly. Returns
    {"reply": str, "tool_calls": [{"name", "input"}, ...], "turns": int}."""
    network_call = network_call or (lambda fn: fn())
    client = _get_client()
    model_name = model_name or os.environ.get("CLAUDE_MODEL", CLAUDE_MODEL)

    os.makedirs(output_dir, exist_ok=True)
    minutes_path = os.path.join(output_dir, "minutes.docx")
    tracker_path = os.path.join(output_dir, "action_tracker.xlsx")

    messages = [{
        "role": "user",
        "content": (
            "Produce meeting minutes and update the action tracker for "
            f"this transcript:\n\n{transcript}"
        ),
    }]
    tool_calls = []

    for turn in range(max_turns):
        try:
            response = network_call(lambda: client.messages.create(
                model=model_name, max_tokens=1024, system=SYSTEM_PROMPT,
                tools=TOOLS, messages=messages,
            ))
        except anthropic.AuthenticationError as exc:
            raise RuntimeError(_NO_AUTH_ERROR) from exc
        except TypeError as exc:
            if "authentication" in str(exc).lower():
                raise RuntimeError(_NO_AUTH_ERROR) from exc
            raise

        messages.append({"role": "assistant", "content": response.content})

        if response.stop_reason != "tool_use":
            final_text = "".join(b.text for b in response.content if b.type == "text").strip()
            return {"reply": final_text, "tool_calls": tool_calls, "turns": turn + 1}

        tool_results = []
        for block in response.content:
            if block.type != "tool_use":
                continue
            tool_calls.append({"name": block.name, "input": block.input})
            if block.name == "write_minutes_docx":
                _write_minutes_docx(minutes_path, block.input)
                result_text = f"Wrote {minutes_path}"
            elif block.name == "write_action_tracker_xlsx":
                _append_action_tracker_xlsx(tracker_path, block.input.get("actions"))
                result_text = f"Updated {tracker_path}"
            else:
                result_text = f"Unknown tool {block.name}"
            tool_results.append({
                "type": "tool_result", "tool_use_id": block.id, "content": result_text,
            })
        messages.append({"role": "user", "content": tool_results})

    return {"reply": "", "tool_calls": tool_calls, "turns": max_turns, "incomplete": True}
